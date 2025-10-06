"""
Word Document Report Generator for Saudi Bank
Creates professional Word documents for case reporting
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
from typing import Dict, Any, List
import os


class WordReportGenerator:
    """Generate Word documents for fraud case reporting"""

    def __init__(self, bank_name="Saudi National Bank", bank_logo_path=None):
        self.bank_name = bank_name
        self.bank_logo_path = bank_logo_path

    def create_case_report(self,
                          transaction: Dict[str, Any],
                          classification: Dict[str, Any],
                          monitoring_review: Dict[str, Any],
                          investigation: Dict[str, Any],
                          output_path: str) -> str:
        """
        Generate comprehensive Word document for a fraud case

        Args:
            transaction: Transaction data
            classification: Initial classification results
            monitoring_review: Monitoring team review
            investigation: Investigation results
            output_path: Path to save the document

        Returns:
            Path to generated document
        """
        doc = Document()

        # Set up document styles
        self._setup_styles(doc)

        # Header with bank info
        self._add_header(doc, "FRAUD INVESTIGATION CASE REPORT")

        # Case summary box
        self._add_case_summary_box(doc, transaction, investigation, monitoring_review)

        # 1. Executive Summary
        self._add_section_header(doc, "1. EXECUTIVE SUMMARY")
        self._add_executive_summary(doc, investigation, transaction)

        # 2. Transaction Details
        self._add_section_header(doc, "2. TRANSACTION DETAILS")
        self._add_transaction_details(doc, transaction)

        # 3. Customer Information
        self._add_section_header(doc, "3. CUSTOMER INFORMATION")
        self._add_customer_info(doc, investigation)

        # 4. Risk Assessment
        self._add_section_header(doc, "4. RISK ASSESSMENT & CLASSIFICATION")
        self._add_risk_assessment(doc, classification, investigation)

        # 5. Investigation Findings
        self._add_section_header(doc, "5. INVESTIGATION FINDINGS")
        self._add_investigation_findings(doc, investigation)

        # 6. SAMA Compliance
        self._add_section_header(doc, "6. SAMA COMPLIANCE & AML REQUIREMENTS")
        self._add_sama_compliance(doc, transaction, investigation)

        # 7. Recommended Actions
        self._add_section_header(doc, "7. RECOMMENDED ACTIONS")
        self._add_recommended_actions(doc, investigation)

        # 8. Approvals Section
        self._add_section_header(doc, "8. APPROVALS & SIGN-OFF")
        self._add_approval_section(doc)

        # Footer
        self._add_footer(doc)

        # Save document
        os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)
        doc.save(output_path)

        return output_path

    def _setup_styles(self, doc):
        """Setup custom styles for the document"""
        styles = doc.styles

        # Title style
        if 'CustomTitle' not in [s.name for s in styles]:
            title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Arial'
            title_style.font.size = Pt(18)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(0, 51, 102)

        # Heading style
        if 'CustomHeading' not in [s.name for s in styles]:
            heading_style = styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
            heading_style.font.name = 'Arial'
            heading_style.font.size = Pt(14)
            heading_style.font.bold = True
            heading_style.font.color.rgb = RGBColor(0, 102, 204)

    def _add_header(self, doc, title):
        """Add document header"""
        # Bank name
        bank_para = doc.add_paragraph()
        bank_run = bank_para.add_run(self.bank_name.upper())
        bank_run.font.size = Pt(16)
        bank_run.font.bold = True
        bank_run.font.color.rgb = RGBColor(0, 51, 102)
        bank_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(14)
        title_run.font.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Divider
        doc.add_paragraph('_' * 80)
        doc.add_paragraph()

    def _add_case_summary_box(self, doc, transaction, investigation, monitoring_review):
        """Add highlighted case summary box"""
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'

        cells = [
            ('Case ID:', monitoring_review.get('case_id', 'N/A')),
            ('Transaction ID:', transaction['transaction_id']),
            ('Customer ID:', transaction['customer_id']),
            ('Case Status:', investigation['case_status']),
            ('Priority Level:', monitoring_review.get('case_priority', 'N/A')),
            ('Report Date:', datetime.now().strftime('%Y-%m-%d %H:%M'))
        ]

        for i, (label, value) in enumerate(cells):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
            table.rows[i].cells[1].text = str(value)

        doc.add_paragraph()

    def _add_section_header(self, doc, title):
        """Add section header"""
        para = doc.add_paragraph()
        run = para.add_run(title)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 102, 204)
        doc.add_paragraph()

    def _add_executive_summary(self, doc, investigation, transaction):
        """Add executive summary section"""
        doc.add_paragraph(
            f"Final Classification: {investigation['final_classification']}\n"
            f"Confidence Level: {investigation['confidence']:.0%}\n"
            f"Transaction Amount: {transaction['amount']:,.2f} {transaction['currency']}\n\n"
            f"Summary:\n{investigation.get('investigator_notes', 'No additional notes')}"
        )
        doc.add_paragraph()

    def _add_transaction_details(self, doc, transaction):
        """Add transaction details table"""
        details = [
            ('Transaction ID', transaction['transaction_id']),
            ('Customer ID', transaction['customer_id']),
            ('Customer Name', transaction.get('customer_name', 'N/A')),
            ('Date & Time', transaction['timestamp']),
            ('Amount', f"{transaction['amount']:,.2f} {transaction['currency']}"),
            ('Beneficiary', transaction.get('beneficiary_name', transaction.get('merchant_name', 'N/A'))),
            ('Beneficiary Bank', transaction.get('beneficiary_bank', 'N/A')),
            ('Beneficiary Country', transaction.get('beneficiary_country', transaction.get('merchant_country', 'N/A'))),
            ('Transfer Type', transaction.get('transfer_type', transaction.get('transaction_type', 'N/A'))),
            ('Transfer Purpose', transaction.get('transfer_purpose', 'N/A')),
            ('ML Fraud Score', f"{transaction['ml_fraud_score']:.3f}"),
            ('Nationality', transaction.get('customer_nationality', 'N/A')),
            ('SAMA AML Flag', 'YES' if transaction.get('sama_aml_flag', False) else 'NO')
        ]

        table = doc.add_table(rows=len(details), cols=2)
        table.style = 'Light List Accent 1'

        for i, (label, value) in enumerate(details):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
            table.rows[i].cells[1].text = str(value)

        doc.add_paragraph()

    def _add_customer_info(self, doc, investigation):
        """Add customer information"""
        doc.add_paragraph(f"Profile Summary: {investigation['customer_profile_summary']}")
        doc.add_paragraph(f"Login Activity: {investigation['login_summary']}")
        doc.add_paragraph(f"Device Information: {investigation['device_summary']}")
        doc.add_paragraph()

    def _add_risk_assessment(self, doc, classification, investigation):
        """Add risk assessment details"""
        doc.add_paragraph(f"Initial Classification: {classification['classification']}")
        doc.add_paragraph(f"Confidence: {classification['confidence']:.0%}\n")

        doc.add_paragraph("Risk Factors Identified:", style='List Bullet')
        for factor in classification.get('risk_factors', []):
            doc.add_paragraph(factor, style='List Bullet 2')

        doc.add_paragraph("\nBehavioral Analysis:")
        behavioral = investigation['behavioral_analysis']
        doc.add_paragraph(f"• Profile Risk: {behavioral['profile_risk']}")
        doc.add_paragraph(f"• Login Risk: {behavioral['login_risk']}")
        doc.add_paragraph(f"• Device Risk: {behavioral['device_risk']}")
        doc.add_paragraph()

    def _add_investigation_findings(self, doc, investigation):
        """Add investigation findings"""
        doc.add_paragraph("Data Sources Analyzed:", style='List Bullet')
        for source in investigation['data_sources_checked']:
            doc.add_paragraph(source.replace('_', ' ').title(), style='List Bullet 2')

        doc.add_paragraph("\nDetailed Analysis:")
        doc.add_paragraph(investigation['investigation_summary'])
        doc.add_paragraph()

        if investigation['behavioral_analysis']['behavioral_anomalies']:
            doc.add_paragraph("Behavioral Anomalies Detected:", style='List Bullet')
            for anomaly in investigation['behavioral_analysis']['behavioral_anomalies']:
                doc.add_paragraph(anomaly, style='List Bullet 2')

        doc.add_paragraph()

    def _add_sama_compliance(self, doc, transaction, investigation):
        """Add SAMA compliance section"""
        doc.add_paragraph("SAMA AML/CFT Compliance Check:\n")

        sama_checks = []

        # Transaction amount threshold
        if transaction['amount'] > 20000:
            sama_checks.append(
                f"✓ Large Transaction Reporting: Amount {transaction['amount']:,.2f} SAR "
                "exceeds SAMA threshold of 20,000 SAR"
            )

        # High-risk country
        risk_countries = ['Iran', 'Yemen', 'Syria', 'North Korea', 'Nigeria']
        beneficiary_country = transaction.get('beneficiary_country', transaction.get('merchant_country', 'Unknown'))
        if beneficiary_country in risk_countries:
            sama_checks.append(
                f"✓ High-Risk Jurisdiction: Transfer to {beneficiary_country} "
                "requires enhanced due diligence"
            )

        # Suspicious activity report
        if investigation['case_status'] in ['CONFIRMED_FRAUD', 'SUSPECTED_FRAUD']:
            sama_checks.append(
                "✓ Suspicious Activity Report (SAR): Case requires filing with SAMA FIU"
            )

        for check in sama_checks:
            doc.add_paragraph(check)

        doc.add_paragraph(
            "\nRegulatory Framework:\n"
            "• Anti-Money Laundering Law (Royal Decree No. M/31)\n"
            "• SAMA AML/CFT Rules 2018\n"
            "• FATF Recommendations Compliance"
        )
        doc.add_paragraph()

    def _add_recommended_actions(self, doc, investigation):
        """Add recommended actions"""
        doc.add_paragraph("Immediate Actions Required:\n", style='List Bullet')
        for i, action in enumerate(investigation['recommended_actions'], 1):
            doc.add_paragraph(f"{i}. {action}", style='List Bullet 2')

        doc.add_paragraph()

    def _add_approval_section(self, doc):
        """Add approval and sign-off section"""
        table = doc.add_table(rows=4, cols=3)
        table.style = 'Light Grid'

        headers = ['Role', 'Name', 'Signature & Date']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True

        roles = ['Fraud Analyst', 'Team Manager', 'Compliance Officer']
        for i, role in enumerate(roles, 1):
            table.rows[i].cells[0].text = role
            table.rows[i].cells[1].text = ''
            table.rows[i].cells[2].text = ''

        doc.add_paragraph()

    def _add_footer(self, doc):
        """Add document footer"""
        doc.add_paragraph('_' * 80)
        footer = doc.add_paragraph()
        footer.add_run(
            f"CONFIDENTIAL - {self.bank_name} - Fraud Investigation Report\n"
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
            "This document contains confidential information and is subject to SAMA regulations."
        )
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            run.font.size = Pt(8)
            run.font.italic = True


if __name__ == "__main__":
    print("Word Report Generator for Saudi Bank - Ready")
    print("Use in conjunction with main.py to generate case reports")
