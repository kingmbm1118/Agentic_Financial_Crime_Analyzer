"""
Word Batch Report Generator for Saudi Bank
Creates summary reports for multiple fraud cases
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from typing import Dict, Any, List
import os


class WordBatchReportGenerator:
    """Generate Word documents for batch fraud case reporting"""

    def __init__(self, bank_name="Saudi National Bank"):
        self.bank_name = bank_name

    def create_batch_report(self,
                           results: List[Dict[str, Any]],
                           statistics: Dict[str, Any],
                           output_path: str) -> str:
        """
        Generate comprehensive batch fraud analysis report

        Args:
            results: List of case processing results
            statistics: Summary statistics
            output_path: Path to save the document

        Returns:
            Path to generated document
        """
        doc = Document()

        # Header
        self._add_header(doc)

        # Executive Summary
        self._add_executive_summary(doc, statistics, results)

        # Statistics Dashboard
        self._add_statistics_dashboard(doc, statistics)

        # High Priority Cases
        self._add_high_priority_cases(doc, results)

        # SAMA Compliance Summary
        self._add_sama_compliance_summary(doc, results)

    def create_closed_cases_report(self,
                                   results: List[Dict[str, Any]],
                                   output_path: str) -> str:
        """
        Generate aggregated report for closed cases (low-risk/legitimate transfers)

        Args:
            results: List of case processing results
            output_path: Path to save the document

        Returns:
            Path to generated document
        """
        doc = Document()

        # Filter closed cases only
        closed_cases = [r for r in results if r.get('monitoring_review', {}).get('action') == 'AGREE_CLOSE']

        # Header
        header = doc.add_heading(f'{self.bank_name}', 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER

        title = doc.add_heading('CLOSED CASES REFERENCE REPORT', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Total Closed Cases: {len(closed_cases)}")
        doc.add_paragraph()

        # Summary
        doc.add_heading('SUMMARY', level=2)
        p = doc.add_paragraph()
        p.add_run(f"Total transfers reviewed and closed: {len(closed_cases)}\n")
        p.add_run(f"All cases classified as LOW RISK or LEGITIMATE\n")
        p.add_run(f"No further investigation required")
        doc.add_paragraph()

        # Closed Cases Table
        doc.add_heading('CLOSED CASES DETAILS', level=2)

        if closed_cases:
            # Create table
            table = doc.add_table(rows=1, cols=7)
            table.style = 'Light List Accent 1'

            # Header row
            headers = ['Transaction ID', 'Customer ID', 'Amount (SAR)', 'Beneficiary Country',
                      'ML Score', 'Classification', 'Reason']
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True

            # Data rows
            for result in closed_cases:
                transaction = result['transaction']
                classification = result['classification']

                row = table.add_row()
                row.cells[0].text = transaction['transaction_id']
                row.cells[1].text = transaction['customer_id']
                row.cells[2].text = f"{transaction['amount']:,.2f}"
                row.cells[3].text = transaction.get('beneficiary_country',
                                                    transaction.get('merchant_country', 'N/A'))
                row.cells[4].text = f"{transaction['ml_fraud_score']:.2f}"
                row.cells[5].text = classification['classification']
                row.cells[6].text = classification['reasoning'][:100] + "..." if len(classification['reasoning']) > 100 else classification['reasoning']
        else:
            doc.add_paragraph("No closed cases in this batch.")

        doc.add_paragraph()

        # Statistics
        doc.add_heading('STATISTICS', level=2)
        if closed_cases:
            avg_amount = sum(r['transaction']['amount'] for r in closed_cases) / len(closed_cases)
            avg_ml_score = sum(r['transaction']['ml_fraud_score'] for r in closed_cases) / len(closed_cases)

            stats_table = doc.add_table(rows=3, cols=2)
            stats_table.style = 'Light Grid Accent 1'

            stats_table.rows[0].cells[0].text = 'Average Transfer Amount'
            stats_table.rows[0].cells[1].text = f"{avg_amount:,.2f} SAR"

            stats_table.rows[1].cells[0].text = 'Average ML Fraud Score'
            stats_table.rows[1].cells[1].text = f"{avg_ml_score:.3f}"

            stats_table.rows[2].cells[0].text = 'Risk Level'
            stats_table.rows[2].cells[1].text = "LOW - All cases cleared"

        # Footer
        doc.add_paragraph()
        doc.add_paragraph("─" * 80)
        footer = doc.add_paragraph()
        footer.add_run("CONFIDENTIAL - For Internal Use Only\n").bold = True
        footer.add_run(f"Generated by {self.bank_name} Fraud Detection System\n")
        footer.add_run("Retention Period: 5 years as per SAMA regulations")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Save
        os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)
        doc.save(output_path)

        return output_path

        # Detailed Case List
        self._add_detailed_case_list(doc, results)

        # Recommendations
        self._add_recommendations(doc, results, statistics)

        # Footer
        self._add_footer(doc)

        # Save
        os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)
        doc.save(output_path)

        return output_path

    def _add_header(self, doc):
        """Add report header"""
        # Bank name
        bank_para = doc.add_paragraph()
        bank_run = bank_para.add_run(self.bank_name.upper())
        bank_run.font.size = Pt(16)
        bank_run.font.bold = True
        bank_run.font.color.rgb = RGBColor(0, 51, 102)
        bank_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run("FRAUD DETECTION BATCH ANALYSIS REPORT")
        title_run.font.size = Pt(14)
        title_run.font.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Report info
        info_para = doc.add_paragraph()
        info_run = info_para.add_run(
            f"Report Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
            f"Reporting Period: Last 7 Days"
        )
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph('_' * 80)
        doc.add_paragraph()

    def _add_executive_summary(self, doc, statistics, results):
        """Add executive summary"""
        # Filter out None and error results
        valid_results = [r for r in results if r and 'error' not in r]

        heading = doc.add_paragraph()
        heading_run = heading.add_run("EXECUTIVE SUMMARY")
        heading_run.font.size = Pt(12)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(0, 102, 204)
        doc.add_paragraph()

        confirmed_fraud = sum(1 for r in valid_results
                             if r.get('investigation', {}).get('case_status') == 'CONFIRMED_FRAUD')

        summary_text = f"""
This report summarizes the analysis of {statistics['total']} transactions processed through the
fraud detection system during the reporting period.

Key Findings:
• {statistics['flagged']} transactions flagged as high risk ({statistics['flagged']/statistics['total']*100:.1f}%)
• {statistics['cases_created']} cases created for investigation
• {confirmed_fraud} cases confirmed as fraudulent activity
• {statistics.get('investigate', 0)} transactions require further monitoring
• Average ML Fraud Score: {statistics.get('avg_ml_score', 0):.3f}

The fraud detection system successfully identified and escalated suspicious transactions
in compliance with SAMA AML/CFT regulations.
"""
        doc.add_paragraph(summary_text)
        doc.add_paragraph()

    def _add_statistics_dashboard(self, doc, statistics):
        """Add statistics in table format"""
        heading = doc.add_paragraph()
        heading_run = heading.add_run("STATISTICS DASHBOARD")
        heading_run.font.size = Pt(12)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(0, 102, 204)
        doc.add_paragraph()

        # Create statistics table
        table = doc.add_table(rows=6, cols=3)
        table.style = 'Light Grid Accent 1'

        # Headers
        headers = ['Metric', 'Count', 'Percentage']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True

        # Data rows
        total = statistics['total']
        data = [
            ('Total Transactions Analyzed', statistics['total'], '100.0%'),
            ('Flagged (High Risk)', statistics['flagged'],
             f"{statistics['flagged']/total*100:.1f}%"),
            ('Needs Investigation', statistics.get('investigate', 0),
             f"{statistics.get('investigate', 0)/total*100:.1f}%"),
            ('Non-Fraud (Low Risk)', statistics.get('non_fraud', 0),
             f"{statistics.get('non_fraud', 0)/total*100:.1f}%"),
            ('Cases Created', statistics['cases_created'],
             f"{statistics['cases_created']/total*100:.1f}%"),
            ('Confirmed Fraud', statistics['confirmed_fraud'],
             f"{statistics['confirmed_fraud']/total*100:.1f}%")
        ]

        for i, (metric, count, pct) in enumerate(data, 1):
            table.rows[i].cells[0].text = metric
            table.rows[i].cells[1].text = str(count)
            table.rows[i].cells[2].text = pct

        doc.add_paragraph()

    def _add_high_priority_cases(self, doc, results):
        """Add high priority cases section"""
        heading = doc.add_paragraph()
        heading_run = heading.add_run("HIGH PRIORITY CASES")
        heading_run.font.size = Pt(12)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(0, 102, 204)
        doc.add_paragraph()

        # Filter high priority cases
        high_priority = [r for r in results
                        if r.get('monitoring_review', {}).get('case_priority') == 'HIGH'][:10]

        if not high_priority:
            doc.add_paragraph("No high priority cases identified in this period.")
        else:
            doc.add_paragraph(
                f"Total High Priority Cases: {len(high_priority)}\n"
                "The following cases require immediate attention:\n"
            )

            # Create table for high priority cases
            table = doc.add_table(rows=len(high_priority)+1, cols=5)
            table.style = 'Light List Accent 1'

            # Headers
            headers = ['Transaction ID', 'Amount (SAR)', 'ML Score', 'Country', 'Status']
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True

            # Data
            for i, result in enumerate(high_priority, 1):
                txn = result['transaction']
                investigation = result.get('investigation', {})

                table.rows[i].cells[0].text = txn['transaction_id']
                table.rows[i].cells[1].text = f"{txn['amount']:,.2f}"
                table.rows[i].cells[2].text = f"{txn['ml_fraud_score']:.3f}"
                table.rows[i].cells[3].text = txn['merchant_country']
                table.rows[i].cells[4].text = investigation.get('case_status', 'PENDING')

        doc.add_paragraph()

    def _add_sama_compliance_summary(self, doc, results):
        """Add SAMA compliance summary"""
        heading = doc.add_paragraph()
        heading_run = heading.add_run("SAMA COMPLIANCE SUMMARY")
        heading_run.font.size = Pt(12)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(0, 102, 204)
        doc.add_paragraph()

        # Count SAMA flagged transactions
        sama_flagged = sum(1 for r in results
                          if r['transaction'].get('sama_aml_flag', False))

        # Count large transactions
        large_txns = sum(1 for r in results
                        if r['transaction']['amount'] > 20000)

        # Count high-risk countries
        risk_countries = ['Iran', 'Yemen', 'Syria', 'North Korea', 'Nigeria']
        high_risk_countries = sum(1 for r in results
                                 if r['transaction']['merchant_country'] in risk_countries)

        # SAR filings needed
        sar_needed = sum(1 for r in results
                        if r.get('investigation', {}).get('case_status') in
                        ['CONFIRMED_FRAUD', 'SUSPECTED_FRAUD'])

        doc.add_paragraph(
            f"SAMA AML/CFT Compliance Metrics:\n\n"
            f"• Transactions with SAMA AML Flags: {sama_flagged}\n"
            f"• Large Transactions (>20,000 SAR): {large_txns}\n"
            f"• High-Risk Country Transactions: {high_risk_countries}\n"
            f"• Suspicious Activity Reports (SAR) Required: {sar_needed}\n\n"
            "All flagged transactions have been processed in accordance with:\n"
            "- Anti-Money Laundering Law (Royal Decree No. M/31)\n"
            "- SAMA AML/CFT Rules 2018\n"
            "- FATF Recommendations"
        )
        doc.add_paragraph()

    def _add_detailed_case_list(self, doc, results):
        """Add detailed case list"""
        heading = doc.add_paragraph()
        heading_run = heading.add_run("DETAILED CASE LIST")
        heading_run.font.size = Pt(12)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(0, 102, 204)
        doc.add_paragraph()

        # Show all cases with investigation
        investigated = [r for r in results if 'investigation' in r]

        if not investigated:
            doc.add_paragraph("No cases created during this period.")
        else:
            for i, result in enumerate(investigated, 1):
                txn = result['transaction']
                investigation = result['investigation']
                case_id = result.get('monitoring_review', {}).get('case_id', 'N/A')

                doc.add_paragraph(
                    f"Case #{i}: {case_id}\n"
                    f"Transaction: {txn['transaction_id']} | "
                    f"Customer: {txn.get('customer_name', 'N/A')} ({txn['customer_id']})\n"
                    f"Amount: {txn['amount']:,.2f} {txn['currency']} | "
                    f"Country: {txn['merchant_country']}\n"
                    f"Status: {investigation['case_status']} | "
                    f"Confidence: {investigation['confidence']:.0%}\n"
                )

        doc.add_paragraph()

    def _add_recommendations(self, doc, results, statistics):
        """Add recommendations section"""
        heading = doc.add_paragraph()
        heading_run = heading.add_run("RECOMMENDATIONS & NEXT STEPS")
        heading_run.font.size = Pt(12)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(0, 102, 204)
        doc.add_paragraph()

        # Generate recommendations based on findings
        recommendations = []

        if statistics['confirmed_fraud'] > 0:
            recommendations.append(
                f"Immediate Action: {statistics['confirmed_fraud']} confirmed fraud cases "
                "require immediate customer contact and account review."
            )

        if statistics['flagged'] / statistics['total'] > 0.2:
            recommendations.append(
                "High fraud rate detected (>20%). Recommend enhanced monitoring "
                "and review of detection thresholds."
            )

        recommendations.extend([
            "File required Suspicious Activity Reports (SAR) with SAMA Financial Intelligence Unit.",
            "Update customer risk profiles based on investigation findings.",
            "Conduct training session for branch staff on identified fraud patterns.",
            "Review and enhance transaction monitoring rules based on new patterns.",
            "Ensure all cases are documented in compliance with SAMA record-keeping requirements (5 years minimum)."
        ])

        for i, rec in enumerate(recommendations, 1):
            doc.add_paragraph(f"{i}. {rec}")

        doc.add_paragraph()

    def _add_footer(self, doc):
        """Add document footer"""
        doc.add_paragraph('_' * 80)
        footer = doc.add_paragraph()
        footer.add_run(
            f"CONFIDENTIAL - {self.bank_name} - Batch Fraud Analysis Report\n"
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
            "This document contains confidential information and is subject to SAMA regulations.\n"
            "Retention Period: 5 years as per SAMA AML/CFT Rules"
        )
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            run.font.size = Pt(8)
            run.font.italic = True


if __name__ == "__main__":
    print("Word Batch Report Generator for Saudi Bank - Ready")
